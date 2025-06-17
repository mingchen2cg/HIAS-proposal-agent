# file_utils.py

import os
import tempfile
import base64
import uuid
import markdown
import fitz  # PyMuPDF
import regex as re
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from docx import Document
from weasyprint import HTML
from config import IMAGE_DIR

def extract_text_and_images_from_pdf(file_path):
    """Extract text using PyPDFLoader and images using PyMuPDF."""
    loader = PyPDFLoader(file_path)
    docs = loader.load()
    extracted_text = "\n\n".join([doc.page_content for doc in docs])
    
    image_paths = []
    try:
        pdf_document = fitz.Document(file_path)
        for page in pdf_document:
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_filename = f"{IMAGE_DIR}/pdf_image_{uuid.uuid4().hex}.png"
                with open(image_filename, "wb") as img_file:
                    img_file.write(image_bytes)
                image_paths.append(image_filename)
        pdf_document.close()
    except Exception as e:
        st.warning(f"Image extraction from PDF failed: {str(e)}. Only text will be extracted.")
    return extracted_text, image_paths

def extract_text_and_images_from_docx(file_path):
    """Extract text using Docx2txtLoader and images using python-docx."""
    loader = Docx2txtLoader(file_path)
    docs = loader.load()
    extracted_text = "\n\n".join([doc.page_content for doc in docs])

    image_paths = []
    try:
        doc = Document(file_path)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_filename = f"{IMAGE_DIR}/docx_image_{uuid.uuid4().hex}.png"
                with open(image_filename, "wb") as img_file:
                    img_file.write(rel.target_part.blob)
                image_paths.append(image_filename)
    except Exception as e:
        st.warning(f"Image extraction from DOCX failed: {str(e)}. Only text will be extracted.")
    return extracted_text, image_paths

def extract_text_from_file(uploaded_file):
    """Extract text and images from an uploaded file based on its type."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.getbuffer())
        file_path = tmp_file.name

    try:
        if file_path.endswith(".pdf"):
            extracted_text, image_paths = extract_text_and_images_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            extracted_text, image_paths = extract_text_and_images_from_docx(file_path)
        else:
            raise ValueError("Only PDF and DOCX files are supported.")
        
        st.session_state.extracted_images.extend(image_paths)
        os.unlink(file_path)
        return extracted_text
    except Exception as e:
        os.unlink(file_path)
        raise e

def convert_md_to_docx(md_text, output_filename, images=None):
    """Convert markdown text to a DOCX document, including images."""
    doc = Document()
    if images:
        for img_path in images[:2]:  # Limit to first 2 images
            try:
                doc.add_picture(img_path, width=2000000)
            except Exception:
                pass

    lines = md_text.split('\n')
    in_mermaid = False
    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith("```mermaid"):
            in_mermaid = True
        elif in_mermaid and stripped_line.startswith("```"):
            in_mermaid = False
            doc.add_paragraph("--- Mermaid Workflow Diagram ---")
        elif not in_mermaid:
            if stripped_line.startswith('# '):
                doc.add_heading(stripped_line[2:], level=1)
            elif stripped_line.startswith('## '):
                doc.add_heading(stripped_line[3:], level=2)
            elif stripped_line.startswith('### '):
                doc.add_heading(stripped_line[4:], level=3)
            elif stripped_line and not stripped_line.startswith('```'):
                doc.add_paragraph(stripped_line)
    
    doc.save(output_filename)
    return output_filename

def convert_md_to_pdf(md_text, output_filename, images=None):
    """Convert markdown text to a PDF document, including images."""
    html_content = markdown.markdown(md_text, extensions=['extra', 'codehilite'])
    
    image_tags = ""
    if images:
        for img_path in images[:2]:  # Limit to first 2 images
            try:
                with open(img_path, "rb") as img_file:
                    img_data = base64.b64encode(img_file.read()).decode()
                img_ext = os.path.splitext(img_path)[1][1:]
                image_tags += f'<img src="data:image/{img_ext};base64,{img_data}" style="max-width: 200px; margin: 10px 0;">'
            except Exception:
                pass

    html_content = re.sub(r'<pre><code>mermaid.*?</code></pre>', '<div class="mermaid-placeholder">--- Mermaid Workflow Diagram ---</div>', html_content, flags=re.DOTALL)
    
    styled_html = f"""
    <html><head><style>
        body {{ font-family: Arial, sans-serif; margin: 2cm; }}
        h1 {{ color: #2c3e50; }} h2 {{ color: #3498db; }} h3 {{ color: #555; }}
        .mermaid-placeholder {{ text-align: center; padding: 20px; border: 1px dashed #ccc; margin: 20px 0; color: #777; font-style: italic; }}
    </style></head><body>{image_tags}{html_content}</body></html>
    """
    HTML(string=styled_html).write_pdf(output_filename)
    return output_filename